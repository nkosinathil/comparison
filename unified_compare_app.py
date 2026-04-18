#!/usr/bin/env python3
"""
Unified Offline File Comparison Application

Creates one comparison application that can compare different categories of files:
- emails (.eml, .msg)
- pdf (.pdf)
- excel (.xlsx, .xlsm)
- word (.docx)
- text (.txt, .csv, .log, .json, .xml, .html, .htm, .md, .rtf)
- tiff (.tif, .tiff)
- all (all supported types)

The report layout follows the same style/pattern as compare_emails.py:
- results.csv
- report.html

Comparison model:
- exact raw hash match
- canonical/content hash match
- attachment hash match (emails only)
- near duplicate (simhash + token Jaccard)
- optional semantic similarity via local Ollama embeddings
- TIFF OCR/text extraction is optional; if OCR libraries are unavailable, TIFF falls back to raw-hash-only comparison.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import math
import re
import sqlite3
import sys
import time
from dataclasses import dataclass, asdict
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# -----------------------------
# Extension groups
# -----------------------------
TEXT_EXTS = {".txt", ".csv", ".log", ".json", ".xml", ".html", ".htm", ".md", ".rtf"}
WORD_EXTS = {".docx"}
EXCEL_EXTS = {".xlsx", ".xlsm"}
PDF_EXTS = {".pdf"}
EMAIL_EXTS = {".eml", ".msg"}
TIFF_EXTS = {".tif", ".tiff"}
SUPPORTED_EXTS = TEXT_EXTS | WORD_EXTS | EXCEL_EXTS | PDF_EXTS | EMAIL_EXTS | TIFF_EXTS

TYPE_MAP = {
    "email": EMAIL_EXTS,
    "pdf": PDF_EXTS,
    "excel": EXCEL_EXTS,
    "word": WORD_EXTS,
    "text": TEXT_EXTS,
    "tiff": TIFF_EXTS,
    "all": SUPPORTED_EXTS,
}

_WORD_RE = re.compile(r"[A-Za-z0-9_']{2,}")


# -----------------------------
# Generic helpers
# -----------------------------
def normalize_whitespace(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def clean_text(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def tokenize(s: str) -> List[str]:
    return [m.group(0).lower() for m in _WORD_RE.finditer(s)]


def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def sha256_text(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8", errors="ignore")).hexdigest()


def token_jaccard(tokens_a: List[str], tokens_b: List[str]) -> float:
    if not tokens_a or not tokens_b:
        return 0.0
    a, b = set(tokens_a), set(tokens_b)
    inter = len(a & b)
    union = len(a | b)
    return float(inter) / float(union) if union else 0.0


def _token_hash64(token: str) -> int:
    h = hashlib.sha256(token.encode("utf-8", errors="ignore")).digest()
    return int.from_bytes(h[:8], "big", signed=False)


def simhash64(tokens: List[str]) -> str:
    if not tokens:
        return "0" * 16
    v = [0] * 64
    for t in tokens:
        h = _token_hash64(t)
        for i in range(64):
            v[i] += 1 if ((h >> i) & 1) else -1
    out = 0
    for i in range(64):
        if v[i] > 0:
            out |= (1 << i)
    return f"{out:016x}"


def simhash_distance(a: int, b: int) -> int:
    return (a ^ b).bit_count()


def _strip_html(html: str) -> str:
    html = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", html)
    html = re.sub(r"(?is)<br\s*/?>", "\n", html)
    html = re.sub(r"(?is)</p\s*>", "\n\n", html)
    html = re.sub(r"(?is)<.*?>", " ", html)
    html = re.sub(r"&nbsp;", " ", html, flags=re.I)
    html = re.sub(r"&amp;", "&", html, flags=re.I)
    html = re.sub(r"&lt;", "<", html, flags=re.I)
    html = re.sub(r"&gt;", ">", html, flags=re.I)
    return normalize_whitespace(html)


def ensure_out_dir(out_dir: Path) -> None:
    out_dir.mkdir(parents=True, exist_ok=True)


def canonical_attachment_set_hash(attachment_hashes: List[str]) -> str:
    if not attachment_hashes:
        return ""
    return sha256_text("|".join(sorted(attachment_hashes)))


# -----------------------------
# Optional semantic similarity
# -----------------------------
def embed_text_ollama(text: str, base_url: str, model: str) -> Optional[List[float]]:
    try:
        import requests  # type: ignore
    except Exception:
        return None
    text = text[:20000]
    try:
        r = requests.post(
            f"{base_url.rstrip('/')}/api/embeddings",
            json={"model": model, "prompt": text},
            timeout=60,
        )
        if r.status_code != 200:
            return None
        data = r.json()
        emb = data.get("embedding")
        if isinstance(emb, list) and emb and isinstance(emb[0], (int, float)):
            return [float(x) for x in emb]
    except Exception:
        return None
    return None


def cosine_similarity(a: List[float], b: List[float]) -> float:
    if not a or not b or len(a) != len(b):
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(y * y for y in b))
    if na == 0.0 or nb == 0.0:
        return 0.0
    return float(dot) / float(na * nb)


def tfidf_cosine_similarity(tokens_a: List[str], tokens_b: List[str]) -> float:
    """Compute TF-IDF weighted cosine similarity between two token lists."""
    if not tokens_a or not tokens_b:
        return 0.0

    from collections import Counter

    tf_a = Counter(tokens_a)
    tf_b = Counter(tokens_b)
    vocab = set(tf_a.keys()) | set(tf_b.keys())
    if not vocab:
        return 0.0

    doc_count = {}
    for w in vocab:
        doc_count[w] = (1 if w in tf_a else 0) + (1 if w in tf_b else 0)

    import math as _m

    def _tfidf(tf: Counter, total: int) -> Dict[str, float]:
        vec: Dict[str, float] = {}
        for w in vocab:
            raw_tf = tf.get(w, 0) / total if total else 0.0
            idf = _m.log(2.0 / doc_count[w]) + 1.0
            vec[w] = raw_tf * idf
        return vec

    va = _tfidf(tf_a, len(tokens_a))
    vb = _tfidf(tf_b, len(tokens_b))

    dot = sum(va[w] * vb[w] for w in vocab)
    na = _m.sqrt(sum(v * v for v in va.values()))
    nb = _m.sqrt(sum(v * v for v in vb.values()))
    if na == 0.0 or nb == 0.0:
        return 0.0
    return float(dot) / float(na * nb)


def combined_comparison_score(
    raw_match: bool,
    canonical_match: bool,
    body_match: bool,
    attachment_overlap_count: int,
    simhash_distance_value: int,
    jaccard_value: float,
    cosine_tfidf_value: float,
    semantic_value: Optional[float],
) -> float:
    """Compute a combined 0-1 comparison score from all available signals."""
    if raw_match:
        return 1.0
    if canonical_match:
        return 0.98
    if body_match:
        return 0.95

    score = 0.0
    simhash_score = max(0.0, 1.0 - (simhash_distance_value / 64.0))
    score += 0.20 * simhash_score
    score += 0.30 * jaccard_value
    score += 0.30 * cosine_tfidf_value

    if semantic_value is not None:
        score = score * 0.6 + semantic_value * 0.4
    else:
        score = score / 0.80

    if attachment_overlap_count > 0:
        score = min(1.0, score + 0.10 * attachment_overlap_count)

    return round(min(1.0, max(0.0, score)), 6)


# -----------------------------
# Parsed record
# -----------------------------
@dataclass
class ParsedRecord:
    path: str
    file_type: str
    subject: str = ""
    sender: str = ""
    recipients: str = ""
    date: str = ""
    body_text_full: str = ""
    body_text_clean: str = ""
    attachment_hashes: Optional[List[str]] = None

    def to_dict(self) -> dict:
        d = asdict(self)
        if d["attachment_hashes"] is None:
            d["attachment_hashes"] = []
        return d


# -----------------------------
# Parsers
# -----------------------------
def _read_text_best_effort(path: Path) -> str:
    data = path.read_bytes()
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="ignore")


def parse_text(path: Path) -> ParsedRecord:
    text = _read_text_best_effort(path)
    return ParsedRecord(
        path=str(path),
        file_type=path.suffix.lower().lstrip(".") or "text",
        body_text_full=text.strip(),
        body_text_clean=clean_text(text),
        attachment_hashes=[],
    )


def parse_docx(path: Path) -> ParsedRecord:
    try:
        import docx
    except Exception as e:
        raise RuntimeError("python-docx not installed; cannot parse .docx") from e

    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    return ParsedRecord(str(path), "docx", body_text_full=text, body_text_clean=clean_text(text), attachment_hashes=[])


def parse_xlsx(path: Path) -> ParsedRecord:
    try:
        import openpyxl
    except Exception as e:
        raise RuntimeError("openpyxl not installed; cannot parse .xlsx/.xlsm") from e

    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    parts: List[str] = []
    for ws in wb.worksheets:
        parts.append(f"== Sheet: {ws.title} ==")
        max_rows = min(ws.max_row or 0, 5000)
        max_cols = min(ws.max_column or 0, 100)
        for r in range(1, max_rows + 1):
            vals: List[str] = []
            for c in range(1, max_cols + 1):
                v = ws.cell(row=r, column=c).value
                if v is not None:
                    vals.append(str(v))
            if vals:
                parts.append(" | ".join(vals))
    text = "\n".join(parts).strip()
    return ParsedRecord(str(path), "xlsx", body_text_full=text, body_text_clean=clean_text(text), attachment_hashes=[])


def parse_pdf(path: Path) -> ParsedRecord:
    text = ""
    try:
        from pypdf import PdfReader  # type: ignore
        reader = PdfReader(str(path))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception:
        try:
            from PyPDF2 import PdfReader  # type: ignore
            reader = PdfReader(str(path))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception as e:
            raise RuntimeError("No PDF text extractor available (install pypdf or PyPDF2)") from e

    text = text.strip()
    return ParsedRecord(str(path), "pdf", body_text_full=text, body_text_clean=clean_text(text), attachment_hashes=[])


def parse_eml(path: Path) -> ParsedRecord:
    msg = BytesParser(policy=policy.default).parsebytes(path.read_bytes())
    subject = str(msg.get("subject", "") or "")
    sender = str(msg.get("from", "") or "")
    recipients = ", ".join([str(x) for x in (msg.get_all("to", []) or [])])
    date = str(msg.get("date", "") or "")

    body_parts: List[str] = []
    attachment_hashes: List[str] = []

    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            disp = part.get_content_disposition()
            if disp == "attachment":
                try:
                    payload = part.get_payload(decode=True) or b""
                    attachment_hashes.append(sha256_bytes(payload))
                except Exception:
                    pass
                continue
            if ctype == "text/plain":
                try:
                    body_parts.append(part.get_content() or "")
                except Exception:
                    pass
            elif ctype == "text/html" and not body_parts:
                try:
                    body_parts.append(_strip_html(part.get_content() or ""))
                except Exception:
                    pass
    else:
        ctype = msg.get_content_type()
        try:
            if ctype == "text/plain":
                body_parts.append(msg.get_content() or "")
            elif ctype == "text/html":
                body_parts.append(_strip_html(msg.get_content() or ""))
        except Exception:
            pass

    body_full = "\n".join([p for p in body_parts if p]).strip()
    return ParsedRecord(
        path=str(path),
        file_type="eml",
        subject=subject,
        sender=sender,
        recipients=recipients,
        date=date,
        body_text_full=body_full,
        body_text_clean=clean_text(body_full),
        attachment_hashes=attachment_hashes,
    )


def parse_msg(path: Path) -> ParsedRecord:
    try:
        import extract_msg  # type: ignore
    except Exception as e:
        raise RuntimeError("extract_msg not installed; cannot parse .msg") from e

    m = extract_msg.Message(str(path))
    m.process()
    subject = m.subject or ""
    sender = m.sender or ""
    recipients = m.to or ""
    date = m.date or ""
    body_full = m.body or ""
    if not body_full and getattr(m, "htmlBody", None):
        body_full = _strip_html(m.htmlBody or "")

    attachment_hashes: List[str] = []
    try:
        for att in (m.attachments or []):
            try:
                attachment_hashes.append(sha256_bytes(att.data or b""))
            except Exception:
                pass
    except Exception:
        pass

    return ParsedRecord(
        path=str(path),
        file_type="msg",
        subject=subject,
        sender=sender,
        recipients=recipients,
        date=date,
        body_text_full=body_full.strip(),
        body_text_clean=clean_text(body_full),
        attachment_hashes=attachment_hashes,
    )


def parse_tiff(path: Path) -> ParsedRecord:
    """
    TIFF parser with graceful fallback.
    Priority:
    1. OCR with Pillow + pytesseract, if available.
    2. No OCR available -> raw-bytes fallback only (empty text).
    """
    text = ""
    try:
        from PIL import Image  # type: ignore
        img = Image.open(str(path))
        try:
            import pytesseract  # type: ignore
            parts: List[str] = []
            frame_idx = 0
            while True:
                try:
                    img.seek(frame_idx)
                except EOFError:
                    break
                try:
                    parts.append(pytesseract.image_to_string(img) or "")
                except Exception:
                    parts.append("")
                frame_idx += 1
            text = "\n".join(p for p in parts if p).strip()
        except Exception:
            text = ""
    except Exception:
        text = ""

    return ParsedRecord(
        path=str(path),
        file_type="tiff",
        body_text_full=text,
        body_text_clean=clean_text(text),
        attachment_hashes=[],
    )


def parse_any(path: Path) -> ParsedRecord:
    ext = path.suffix.lower()
    if ext == ".eml":
        return parse_eml(path)
    if ext == ".msg":
        return parse_msg(path)
    if ext in PDF_EXTS:
        return parse_pdf(path)
    if ext in WORD_EXTS:
        return parse_docx(path)
    if ext in EXCEL_EXTS:
        return parse_xlsx(path)
    if ext in TEXT_EXTS:
        return parse_text(path)
    if ext in TIFF_EXTS:
        return parse_tiff(path)
    raise RuntimeError(f"Unsupported extension: {ext}")


# -----------------------------
# Fingerprints / cache
# -----------------------------
def compute_fingerprints(rec: ParsedRecord, raw_bytes: bytes) -> Dict[str, object]:
    body_clean = rec.body_text_clean or ""
    tokens = tokenize(body_clean)

    if rec.file_type in ("eml", "msg"):
        canon = "\n".join([
            f"subject:{normalize_whitespace(rec.subject.lower())}",
            f"from:{normalize_whitespace(rec.sender.lower())}",
            f"to:{normalize_whitespace(rec.recipients.lower())}",
            f"date:{normalize_whitespace(rec.date.lower())}",
            "",
            body_clean,
        ])
    else:
        canon = body_clean

    return {
        "sha256_raw": sha256_bytes(raw_bytes),
        "sha256_canonical": sha256_text(canon),
        "sha256_body_clean": sha256_text(body_clean),
        "simhash64": simhash64(tokens),
        "tokens": tokens[:50000],
        "attachment_hashes": list(rec.attachment_hashes or []),
    }


def open_cache(db_path: Path) -> sqlite3.Connection:
    conn = sqlite3.connect(str(db_path))
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS compare_fingerprints (
            path TEXT PRIMARY KEY,
            mtime REAL NOT NULL,
            raw_sha256 TEXT NOT NULL,
            data_json TEXT NOT NULL
        )
        """
    )
    conn.commit()
    return conn


def cache_get(conn: sqlite3.Connection, path: Path) -> Optional[dict]:
    st = path.stat()
    row = conn.execute(
        "SELECT mtime, raw_sha256, data_json FROM compare_fingerprints WHERE path = ?",
        (str(path),),
    ).fetchone()
    if not row:
        return None
    mtime, raw_sha256, data_json = row
    if float(mtime) != float(st.st_mtime):
        return None
    try:
        if sha256_bytes(path.read_bytes()) != raw_sha256:
            return None
    except Exception:
        return None
    return json.loads(data_json)


def cache_put(conn: sqlite3.Connection, path: Path, raw_sha256: str, data: dict) -> None:
    st = path.stat()
    conn.execute(
        "INSERT OR REPLACE INTO compare_fingerprints (path, mtime, raw_sha256, data_json) VALUES (?, ?, ?, ?)",
        (str(path), float(st.st_mtime), raw_sha256, json.dumps(data, ensure_ascii=False)),
    )


def compute_or_load(path: Path, conn: Optional[sqlite3.Connection]) -> dict:
    if conn is not None:
        cached = cache_get(conn, path)
        if cached is not None:
            return cached

    raw = path.read_bytes()
    rec = parse_any(path)
    fps = compute_fingerprints(rec, raw)
    data = {"record": rec.to_dict(), "fps": fps}

    if conn is not None:
        cache_put(conn, path, fps["sha256_raw"], data)
    return data


# -----------------------------
# Type filtering
# -----------------------------
def selected_extensions(compare_types: List[str]) -> set:
    exts = set()
    for t in compare_types:
        t = t.strip().lower()
        if t not in TYPE_MAP:
            raise ValueError(f"Unsupported compare type: {t}")
        exts |= TYPE_MAP[t]
    return exts


def iter_supported_files(root: Path, exts: set) -> List[Path]:
    files: List[Path] = []
    for p in root.rglob("*"):
        if p.is_file() and p.suffix.lower() in exts:
            files.append(p)
    return files


def source_type_matches(path: Path, allowed_exts: set) -> bool:
    return path.suffix.lower() in allowed_exts


# -----------------------------
# Verdict logic
# -----------------------------
def verdict_for(
    source: dict,
    target: dict,
    simhash_max_dist: int,
    jaccard_near_dup: float,
    semantic_threshold: float,
    semantic_review_threshold: float,
    use_semantic: bool,
    source_emb: Optional[List[float]],
    target_emb: Optional[List[float]],
    cosine_near_dup: float = 0.85,
) -> Tuple[str, List[str], Dict[str, float]]:
    s_fps = source["fps"]
    t_fps = target["fps"]
    reasons: List[str] = []
    scores: Dict[str, float] = {}

    raw_match = s_fps["sha256_raw"] == t_fps["sha256_raw"]
    canonical_match = s_fps["sha256_canonical"] == t_fps["sha256_canonical"]
    body_match = s_fps["sha256_body_clean"] == t_fps["sha256_body_clean"]

    s_attach = set(s_fps.get("attachment_hashes", []) or [])
    t_attach = set(t_fps.get("attachment_hashes", []) or [])
    attachment_overlap = sorted(s_attach.intersection(t_attach))
    attachment_set_match = (
        canonical_attachment_set_hash(list(s_attach)) == canonical_attachment_set_hash(list(t_attach))
        and len(s_attach) > 0
        and len(t_attach) > 0
    )

    try:
        sh_dist = simhash_distance(int(s_fps["simhash64"], 16), int(t_fps["simhash64"], 16))
    except Exception:
        sh_dist = 64
    scores["simhash_distance"] = float(sh_dist)

    source_tokens = s_fps.get("tokens", []) or []
    target_tokens = t_fps.get("tokens", []) or []

    jac = token_jaccard(source_tokens, target_tokens)
    scores["token_jaccard"] = float(jac)

    cosine_tfidf = tfidf_cosine_similarity(source_tokens, target_tokens)
    scores["cosine_tfidf"] = float(cosine_tfidf)

    semantic = None
    if use_semantic and source_emb and target_emb:
        semantic = cosine_similarity(source_emb, target_emb)
        scores["semantic_cosine"] = float(semantic)

    comparison_score = combined_comparison_score(
        raw_match=raw_match,
        canonical_match=canonical_match,
        body_match=body_match,
        attachment_overlap_count=len(attachment_overlap),
        simhash_distance_value=sh_dist,
        jaccard_value=jac,
        cosine_tfidf_value=cosine_tfidf,
        semantic_value=semantic,
    )
    scores["comparison_score"] = float(comparison_score)

    if raw_match or canonical_match:
        if raw_match:
            reasons.append("raw_sha256_match")
        if canonical_match:
            reasons.append("canonical_sha256_match")
        return "IDENTICAL", reasons, scores

    if body_match:
        reasons.append("clean_content_sha256_match")
        return "CONTENT_DUPLICATE", reasons, scores

    if attachment_set_match:
        reasons.append("attachment_set_match")
        return "ATTACHMENT_MATCH", reasons, scores

    if attachment_overlap:
        reasons.append(f"attachment_hash_overlap:{len(attachment_overlap)}")
        return "ATTACHMENT_MATCH", reasons, scores

    near_duplicate_hit = False
    if sh_dist <= simhash_max_dist:
        reasons.append(f"simhash_distance<= {simhash_max_dist} ({sh_dist})")
        near_duplicate_hit = True
    if jac >= jaccard_near_dup:
        reasons.append(f"token_jaccard>= {jaccard_near_dup:.2f} ({jac:.4f})")
        near_duplicate_hit = True
    if cosine_tfidf >= cosine_near_dup:
        reasons.append(f"cosine_tfidf>= {cosine_near_dup:.2f} ({cosine_tfidf:.4f})")
        near_duplicate_hit = True
    if near_duplicate_hit:
        return "NEAR_DUPLICATE", reasons, scores

    if semantic is not None:
        if semantic >= semantic_threshold:
            reasons.append(f"semantic_cosine>= {semantic_threshold:.2f} ({semantic:.4f})")
            return "SEMANTICALLY_SIMILAR", reasons, scores
        if semantic >= semantic_review_threshold:
            reasons.append(f"semantic_cosine_review>= {semantic_review_threshold:.2f} ({semantic:.4f})")
            return "REVIEW_SEMANTIC", reasons, scores

    return "UNRELATED", reasons, scores



# -----------------------------
# HTML report (same style/pattern as compare_emails.py)
# -----------------------------
def write_html_report(
    html_path: Path,
    source_path: str,
    target_root: str,
    rows: List[dict],
    generated_ts: str,
) -> None:
    def esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    headers = list(rows[0].keys()) if rows else []
    table_rows = []
    for r in rows:
        verdict = str(r.get("verdict", ""))
        row_class = ""
        if verdict == "IDENTICAL":
            row_class = " class=\"good\""
        elif verdict in {"NEAR_DUPLICATE", "SEMANTICALLY_SIMILAR", "REVIEW_SEMANTIC", "CONTENT_DUPLICATE", "ATTACHMENT_MATCH"}:
            row_class = " class=\"warn\""
        elif verdict == "UNRELATED":
            row_class = " class=\"neutral\""
        elif verdict == "ERROR":
            row_class = " class=\"bad\""
        tds = "".join(f"<td>{esc(str(r.get(h, '')))}</td>" for h in headers)
        table_rows.append(f"<tr{row_class}>{tds}</tr>")

    html = f"""<!doctype html>
<html>
<head>
<meta charset=\"utf-8\"/>
<title>Comparison Report</title>
<style>
body {{ font-family: Arial, sans-serif; margin: 16px; }}
.small {{ color: #555; font-size: 12px; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #ddd; padding: 6px 8px; font-size: 12px; vertical-align: top; }}
th {{ cursor: pointer; background: #f6f6f6; position: sticky; top: 0; }}
tr:hover {{ background: #fafafa; }}
.good {{ background: #ecfff2; }}
.warn {{ background: #fff7e6; }}
.bad {{ background: #ffecec; }}
.neutral {{ background: #ffffff; }}
</style>
</head>
<body>
<h2>Offline Comparison Report</h2>
<div class=\"small\">
<b>Generated:</b> {esc(generated_ts)}<br/>
<b>Source:</b> {esc(source_path)}<br/>
<b>Targets:</b> {esc(target_root)}<br/>
<b>Rows:</b> {len(rows)}<br/>
<b>Ordering:</b> Highest comparison_score first, with identical files at the top.
</div>
<br/>
<table id=\"t\">
<thead><tr>
{''.join(f'<th onclick="sortTable({i})">{esc(h)}</th>' for i, h in enumerate(headers))}
</tr></thead>
<tbody>
{''.join(table_rows)}
</tbody>
</table>
<script>
function sortTable(n) {{
  var table = document.getElementById("t");
  var tbody = table.tBodies[0];
  var rows = Array.prototype.slice.call(tbody.rows, 0);
  var asc = table.getAttribute("data-sort-col") != n || table.getAttribute("data-sort-dir") != "asc";
  rows.sort(function(a, b) {{
    var x = a.cells[n].innerText || "";
    var y = b.cells[n].innerText || "";
    var nx = parseFloat(x); var ny = parseFloat(y);
    if (!isNaN(nx) && !isNaN(ny)) {{ return asc ? (nx - ny) : (ny - nx); }}
    x = x.toLowerCase(); y = y.toLowerCase();
    if (x < y) return asc ? -1 : 1;
    if (x > y) return asc ? 1 : -1;
    return 0;
  }});
  rows.forEach(function(r) {{ tbody.appendChild(r); }});
  table.setAttribute("data-sort-col", n);
  table.setAttribute("data-sort-dir", asc ? "asc" : "desc");
}}
</script>
</body>
</html>
"""
    html_path.write_text(html, encoding="utf-8")



# -----------------------------
# Main
# -----------------------------
def main() -> int:
    ap = argparse.ArgumentParser(description="Unified Offline File Comparison Application")
    ap.add_argument("--source", required=True, help="Path to source file")
    ap.add_argument("--targets", required=True, help="Folder containing target files (recursively scanned)")
    ap.add_argument("--out", required=True, help="Output folder")
    ap.add_argument("--cache", default="", help="Optional SQLite cache path (e.g., out/fingerprints.sqlite)")
    ap.add_argument(
        "--compare-types",
        default="all",
        help="Comma-separated comparison types: email,pdf,excel,word,text,tiff,all",
    )

    ap.add_argument("--semantic", action="store_true", help="Enable semantic similarity using local Ollama embeddings")
    ap.add_argument("--ollama-url", default="http://localhost:11434", help="Ollama base URL")
    ap.add_argument("--ollama-model", default="nomic-embed-text", help="Ollama embedding model name")

    ap.add_argument("--simhash-max-dist", type=int, default=6, help="SimHash distance threshold for NEAR_DUPLICATE")
    ap.add_argument("--jaccard-near-dup", type=float, default=0.80, help="Token Jaccard threshold for NEAR_DUPLICATE")
    ap.add_argument("--cosine-near-dup", type=float, default=0.85, help="Cosine (TF-IDF) threshold for NEAR_DUPLICATE")
    ap.add_argument("--semantic-threshold", type=float, default=0.90, help="Cosine threshold for SEMANTICALLY_SIMILAR")
    ap.add_argument("--semantic-review-threshold", type=float, default=0.80, help="Cosine threshold for REVIEW_SEMANTIC")
    args = ap.parse_args()

    src_path = Path(args.source).expanduser().resolve()
    tgt_root = Path(args.targets).expanduser().resolve()
    out_dir = Path(args.out).expanduser().resolve()
    ensure_out_dir(out_dir)

    if not src_path.exists() or not src_path.is_file():
        print(f"ERROR: source does not exist: {src_path}", file=sys.stderr)
        return 2
    if not tgt_root.exists() or not tgt_root.is_dir():
        print(f"ERROR: targets folder does not exist: {tgt_root}", file=sys.stderr)
        return 2

    try:
        compare_types = [x.strip().lower() for x in args.compare_types.split(",") if x.strip()]
        allowed_exts = selected_extensions(compare_types)
    except ValueError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 2

    if not source_type_matches(src_path, allowed_exts):
        print(
            f"ERROR: source type {src_path.suffix.lower()} does not match selected compare types: {','.join(compare_types)}",
            file=sys.stderr,
        )
        return 2

    cache_conn = None
    if args.cache:
        cache_path = Path(args.cache).expanduser().resolve()
        cache_path.parent.mkdir(parents=True, exist_ok=True)
        cache_conn = open_cache(cache_path)

    t0 = time.time()

    try:
        source = compute_or_load(src_path, cache_conn)
    except Exception as e:
        print(f"ERROR: failed to parse source: {e}", file=sys.stderr)
        return 3

    target_files = iter_supported_files(tgt_root, allowed_exts)
    if not target_files:
        print(f"No supported target files found in {tgt_root} for compare types: {','.join(compare_types)}")
        return 0

    source_emb = None
    if args.semantic:
        src_text = source["record"].get("body_text_clean", "") or source["record"].get("body_text_full", "")
        source_emb = embed_text_ollama(src_text, base_url=args.ollama_url, model=args.ollama_model)
        if source_emb is None:
            print("[warn] Semantic enabled but Ollama embedding failed. Continuing without semantic.")

    rows: List[dict] = []

    for i, tp in enumerate(target_files, start=1):
        try:
            target = compute_or_load(tp, cache_conn)
        except Exception as e:
            rows.append(
                {
                    "target_path": str(tp),
                    "target_type": tp.suffix.lower().lstrip("."),
                    "target_subject": "",
                    "verdict": "ERROR",
                    "reasons": f"parse_error:{e}",
                    "raw_hash_match": "",
                    "canonical_hash_match": "",
                    "body_hash_match": "",
                    "attachment_match_count": "",
                    "simhash_distance": "",
                    "token_jaccard": "",
                    "cosine_tfidf": "",
                    "semantic_cosine": "",
                    "comparison_score": "",
                }
            )
            continue

        target_emb = None
        semantic_used = False
        if args.semantic and source_emb is not None:
            tgt_text = target["record"].get("body_text_clean", "") or target["record"].get("body_text_full", "")
            target_emb = embed_text_ollama(tgt_text, base_url=args.ollama_url, model=args.ollama_model)
            semantic_used = target_emb is not None

        verdict, reasons, scores = verdict_for(
            source,
            target,
            simhash_max_dist=args.simhash_max_dist,
            jaccard_near_dup=args.jaccard_near_dup,
            semantic_threshold=args.semantic_threshold,
            semantic_review_threshold=args.semantic_review_threshold,
            use_semantic=(args.semantic and source_emb is not None and semantic_used),
            source_emb=source_emb,
            target_emb=target_emb,
            cosine_near_dup=args.cosine_near_dup,
        )

        s_fps = source["fps"]
        t_fps = target["fps"]
        s_attach = set(s_fps.get("attachment_hashes", []) or [])
        t_attach = set(t_fps.get("attachment_hashes", []) or [])
        attach_overlap = sorted(s_attach.intersection(t_attach))

        rows.append(
            {
                "target_path": str(tp),
                "target_type": target["record"].get("file_type", tp.suffix.lower().lstrip(".")),
                "target_subject": target["record"].get("subject", ""),
                "verdict": verdict,
                "reasons": ";".join(reasons),
                "raw_hash_match": int(s_fps["sha256_raw"] == t_fps["sha256_raw"]),
                "canonical_hash_match": int(s_fps["sha256_canonical"] == t_fps["sha256_canonical"]),
                "body_hash_match": int(s_fps["sha256_body_clean"] == t_fps["sha256_body_clean"]),
                "attachment_match_count": len(attach_overlap),
                "simhash_distance": scores.get("simhash_distance", ""),
                "token_jaccard": scores.get("token_jaccard", ""),
                "cosine_tfidf": scores.get("cosine_tfidf", ""),
                "semantic_cosine": scores.get("semantic_cosine", ""),
                "comparison_score": scores.get("comparison_score", ""),
            }
        )

        if i % 50 == 0:
            print(f"Processed {i}/{len(target_files)}...")

    if cache_conn is not None:
        cache_conn.commit()
        cache_conn.close()

    rows.sort(key=lambda r: float(r.get("comparison_score") or 0.0), reverse=True)

    csv_path = out_dir / "results.csv"
    fieldnames = [
        "target_path",
        "target_type",
        "target_subject",
        "verdict",
        "reasons",
        "raw_hash_match",
        "canonical_hash_match",
        "body_hash_match",
        "attachment_match_count",
        "simhash_distance",
        "token_jaccard",
        "cosine_tfidf",
        "semantic_cosine",
        "comparison_score",
    ]
    with csv_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for r in rows:
            writer.writerow(r)

    html_path = out_dir / "report.html"
    write_html_report(
        html_path=html_path,
        source_path=str(src_path),
        target_root=str(tgt_root),
        rows=rows,
        generated_ts=time.strftime("%Y-%m-%d %H:%M:%S"),
    )

    dt = time.time() - t0
    print(f"Done. Compared {len(rows)} targets in {dt:.1f}s")
    print(f"CSV:  {csv_path}")
    print(f"HTML: {html_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
